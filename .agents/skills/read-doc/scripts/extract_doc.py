#!/usr/bin/env python3
"""
Document text extractor for read-doc skill.
Supports: .hwp, .hwpx, .pdf, .docx, .xlsx, .pptx

Usage:
    python extract_doc.py <file-path>
    python extract_doc.py /path/to/doc.hwp
    python extract_doc.py /path/to/report.pdf
"""

import sys
import os

def extract_hwp(path):
    import subprocess, tempfile, html, re
    with tempfile.TemporaryDirectory() as tmp:
        out = os.path.join(tmp, 'out')
        r = subprocess.run(['hwp5html', path, '--output', out], capture_output=True)
        xhtml = os.path.join(out, 'index.xhtml')
        if not os.path.exists(xhtml):
            print("ERROR: hwp5html 실패. pyhwp 설치 확인: pip install pyhwp", file=sys.stderr)
            sys.exit(1)
        with open(xhtml, encoding='utf-8', errors='ignore') as f:
            c = f.read()
    c = re.sub(r'<style[^>]*>.*?</style>', '', c, flags=re.DOTALL)
    c = re.sub(r'<script[^>]*>.*?</script>', '', c, flags=re.DOTALL)
    t = re.sub(r'<[^>]+>', ' ', c)
    t = html.unescape(t)
    t = re.sub(r'[ \t]+', ' ', t)
    t = re.sub(r'\n{3,}', '\n\n', t).strip()
    return t

def extract_hwpx(path):
    import zipfile, re, xml.etree.ElementTree as ET
    with zipfile.ZipFile(path) as z:
        names = z.namelist()
        if 'Preview/PrvText.txt' in names:
            t = z.read('Preview/PrvText.txt').decode('utf-8', errors='ignore')
            for pat in ['</home/', '</var/']:
                i = t.find(pat)
                if i > 0: t = t[:i]
            return t.strip()
        else:
            sections = sorted(
                [n for n in names if re.match(r'Contents/section\d+\.xml', n)],
                key=lambda x: int(re.search(r'section(\d+)', x).group(1))
            )
            texts = []
            for sec in sections:
                root = ET.fromstring(z.read(sec).decode('utf-8', errors='ignore'))
                for node in root.iter('{http://www.hancom.co.kr/hwpml/2012/paragraph}t'):
                    if node.text: texts.append(node.text)
            return '\n'.join(texts)

def extract_pdf(path):
    import subprocess
    r = subprocess.run(['pdftotext', '-layout', path, '-'], capture_output=True, text=True)
    if r.returncode != 0:
        print("ERROR: pdftotext 실패. poppler-utils 설치 확인: sudo dnf install poppler-utils", file=sys.stderr)
        sys.exit(1)
    return r.stdout

def extract_docx(path):
    try:
        import docx
    except ImportError:
        print("ERROR: python-docx 미설치. pip install python-docx", file=sys.stderr)
        sys.exit(1)
    doc = docx.Document(path)
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells: texts.append(' | '.join(cells))
    return '\n'.join(texts)

def extract_xlsx(path):
    try:
        import openpyxl
    except ImportError:
        print("ERROR: openpyxl 미설치. pip install openpyxl", file=sys.stderr)
        sys.exit(1)
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    lines = []
    for name in wb.sheetnames:
        ws = wb[name]
        lines.append(f'=== Sheet: {name} ===')
        for row in ws.iter_rows(values_only=True):
            vals = [str(v) for v in row if v is not None]
            if vals: lines.append('\t'.join(vals))
    return '\n'.join(lines)

def extract_pptx(path):
    import zipfile, re, xml.etree.ElementTree as ET
    NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    with zipfile.ZipFile(path) as z:
        slides = sorted(
            [n for n in z.namelist() if re.match(r'ppt/slides/slide\d+\.xml', n)],
            key=lambda x: int(re.search(r'slide(\d+)', x).group(1))
        )
        lines = []
        for s in slides:
            root = ET.fromstring(z.read(s))
            texts = [t.text for t in root.iter(f'{{{NS}}}t') if t.text]
            if texts: lines.append(' '.join(texts))
    return '\n'.join(lines)

def main():
    if len(sys.argv) < 2:
        print("Usage: python extract_doc.py <file-path>", file=sys.stderr)
        sys.exit(1)

    path = sys.argv[1]
    if not os.path.exists(path):
        print(f"ERROR: 파일을 찾을 수 없습니다: {path}", file=sys.stderr)
        sys.exit(1)

    ext = os.path.splitext(path)[1].lower()
    extractors = {
        '.hwp': extract_hwp,
        '.hwpx': extract_hwpx,
        '.pdf': extract_pdf,
        '.docx': extract_docx,
        '.xlsx': extract_xlsx,
        '.pptx': extract_pptx,
    }

    if ext not in extractors:
        print(f"ERROR: 지원하지 않는 형식: {ext}", file=sys.stderr)
        print(f"지원 형식: {', '.join(extractors.keys())}", file=sys.stderr)
        sys.exit(1)

    result = extractors[ext](path)
    print(result)

if __name__ == '__main__':
    main()
